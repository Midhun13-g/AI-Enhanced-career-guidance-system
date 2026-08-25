import os
import logging
import shutil
from typing import Tuple

logger = logging.getLogger(__name__)

class DoclingOCRProcessor:
    def __init__(self):
        self.docling_available = False
        self.rapidocr_available = False
        
        # Check Docling availability
        try:
            from docling.document_converter import DocumentConverter
            if os.name == "nt" and shutil.which("cl") is None:
                logger.info("Docling disabled on Windows because the MSVC compiler (cl) is unavailable.")
            else:
                self.converter = DocumentConverter()
                self.docling_available = True
                logger.info("Docling DocumentConverter initialized successfully.")
        except Exception as e:
            logger.info(f"Docling not available, using pdfplumber/python-docx fallbacks: {e}")
        
        # Check RapidOCR availability
        try:
            from rapidocr_onnxruntime import RapidOCR
            self.ocr_engine = RapidOCR()
            self.rapidocr_available = True
            logger.info("RapidOCR initialized successfully.")
        except Exception as e:
            logger.info(f"RapidOCR not available: {e}")

    def process_document(self, file_path: str) -> Tuple[str, str]:
        """
        Extract text from file_path (PDF, DOCX, DOC).
        Returns (extracted_text, extraction_method).
        """
        ext = os.path.splitext(file_path)[1].lower()
        extracted_text = ""
        method_used = "fallback"

        # Try Docling first if available
        if self.docling_available:
            try:
                res = self.converter.convert(file_path)
                doc_text = res.document.export_to_markdown()
                if doc_text and len(doc_text.strip()) > 50:
                    return doc_text, "docling"
            except Exception as e:
                logger.warning(f"Docling conversion failed for {file_path}: {e}")

        # Processing PDF
        if ext == ".pdf":
            extracted_text, method_used = self._process_pdf(file_path)
        # Processing DOCX / DOC
        elif ext in [".docx", ".doc"]:
            extracted_text, method_used = self._process_docx(file_path)
        elif ext == ".txt":
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    extracted_text = f.read()
                method_used = "txt_reader"
            except Exception as e:
                logger.error(f"Error reading TXT file: {e}")

        # OCR fallback if text is sparse and OCR is available
        if len(extracted_text.strip()) < 50 and self.rapidocr_available and ext == ".pdf":
            ocr_text = self._process_pdf_ocr(file_path)
            if len(ocr_text.strip()) > len(extracted_text.strip()):
                extracted_text = ocr_text
                method_used = "rapid_ocr"

        return extracted_text, method_used

    def _process_pdf(self, file_path: str) -> Tuple[str, str]:
        text = ""
        # Try pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                pages_text = [page.extract_text() or "" for page in pdf.pages]
                text = "\n".join(pages_text)
                if len(text.strip()) > 50:
                    return text, "pdfplumber"
        except Exception as e:
            logger.info(f"pdfplumber failed: {e}")

        # Try pypdf / PyPDF2
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            pages_text = [page.extract_text() or "" for page in reader.pages]
            text = "\n".join(pages_text)
            if len(text.strip()) > 50:
                return text, "pypdf"
        except Exception as e:
            logger.info(f"pypdf failed: {e}")

        return text, "pdf_fallback"

    def _process_docx(self, file_path: str) -> Tuple[str, str]:
        try:
            import docx
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text:
                    full_text.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))
            text = "\n".join(full_text)
            return text, "python_docx"
        except Exception as e:
            logger.error(f"Error processing DOCX with python-docx: {e}")
            return "", "docx_error"

    def _process_pdf_ocr(self, file_path: str) -> str:
        text_lines = []
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            for page in doc:
                pix = page.get_pixmap()
                img_bytes = pix.tobytes("png")
                ocr_res, _ = self.ocr_engine(img_bytes)
                if ocr_res:
                    for line in ocr_res:
                        if len(line) >= 2 and (isinstance(line[1], tuple) or isinstance(line[1], list)):
                            text_lines.append(line[1][0])
                        elif len(line) >= 2 and isinstance(line[1], str):
                            text_lines.append(line[1])
        except Exception as e:
            logger.warning(f"PyMuPDF OCR extraction failed: {e}")
        return "\n".join(text_lines)

# Singleton document processor instance
doc_processor = DoclingOCRProcessor()
